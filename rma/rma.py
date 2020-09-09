"""
Automated 900 doc generation script.
"""
from . import vendors as v
from docx import Document
from datetime import datetime as dt
from common import common

def get_name():
    return common.set_username("Type your name as it shall appear on the completed forms")

def parse_date(text_date):
    try:
        return dt.strptime(text_date, "%m/%d/%Y").strftime("%Y-%m-%d")
    except:
        return ''

def batch_form_generate():
    username = get_name()
    
    with open('rma/refurb_data.txt', 'r') as rows:
        for row in rows:
            cells = row.split('\t')
            rma = cells[0]
            items = cells[1]
            vendor = cells[2]
            recovered_platform = cells[3]
            return_doc = cells[5]
            form_number = return_doc[11:]
            departure_date = parse_date(cells[6].strip())


            doc = Document('templates/3305-00900-00000.docx')
            table = doc.tables[0]

            cell = table.rows[1].cells[0]
            cell.paragraphs[1].text = cell.paragraphs[1].text.replace('name', username)

            cell = table.rows[1].cells[1]
            cell.paragraphs[1].text = cell.paragraphs[1].text.replace('form_date', dt.today().strftime("%Y-%m-%d"))

            cell = table.rows[1].cells[5]
            cell.paragraphs[1].text = cell.paragraphs[1].text.replace('form_number', form_number)

            cell = table.rows[2].cells[0]
            cell.paragraphs[1].text = cell.paragraphs[1].text.replace('departure_date', departure_date)
            #cell.paragraphs[1].text = "spam!"

            cell = table.rows[3].cells[0]
            cell.paragraphs[1].text = cell.paragraphs[1].text.replace('rma', rma)

            cell = table.rows[3].cells[2]
            cell.paragraphs[1].text = v.contacts[vendor]

            cell = table.rows[3].cells[4]
            cell.paragraphs[1].text = cell.paragraphs[1].text.replace('vendor', vendor)


            table = doc.tables[2]

            cell = table.rows[1].cells[0]
            cell.paragraphs[3].text = items

            table = doc.tables[4]

            cell = table.rows[1].cells[0]
            cell.paragraphs[3].text = cell.paragraphs[3].text.replace('recovered_platform', recovered_platform)




            # Doc Title and Author properties...
            doc.core_properties.title = ("RMA_%s_<Class>-<Series>_%s_Shipping"
                                     % (rma, dt.today().strftime("%Y-%m-%d")))
            author_inits = username.split()
            author_inits[:-1] = [init[0] + '.' for init in author_inits[:-1]]
            doc.core_properties.author = ' '.join(author_inits)

            # Save doc...
            doc.save('save/3305-00900-%s.docx' % form_number)

def main():
    while True:
        header = ''.join(("\n", "-" * 19, "RMA AND SHIPPING MENU", "-" * 19))
        proclist = ["Display HELP file for instructions.", "I know what I'm doing. Generate the docs!"]
        try:
            proc_id = int(common.dynamicmenu_get("Select an action", proclist, header=header))
        except TypeError:
            break
        if proc_id == 0:
            print("\nSpam!")
        elif proc_id == 1:
            batch_form_generate()
            print("Process complete. Check documents for errors.")
        