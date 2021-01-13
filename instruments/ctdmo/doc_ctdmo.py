# Exporting CTDMO procedure results as a MS Word document

from docx import Document

def qct_to_doc(instrument, qct):
    # Open the results template...
    doc = Document('templates/3305-00101-00000.docx')

    # Calling doc.tables will return all tables in the template, where all
    # text will go. Access each of these by index number. Start with the
    # first one, at index 0...
    table = doc.tables[0]

    # partnumber and serialnumber are found in the third row, first cell...
    cell = table.rows[2].cells[0]

    # Since each line is technically a separate paragraph, fill them in
    # one at a time to better preserve formatting...
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('partnumber', instrument.part_no)
    cell.paragraphs[3].text = cell.paragraphs[3].text.replace('serialnumber', instrument.serialnumber)

    # seriesletter is found in the third row, second cell...
    cell = table.rows[2].cells[1]
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('seriesletter', instrument.seriesletter)

    # formnumber is found in the third row, last cell...
    cell = table.rows[2].cells[-1]
    cell.paragraphs[2].text = cell.paragraphs[2].text.replace('formnumber', qct.header['formnumber'])

    # On to the next table...
    table = doc.tables[1]
    cell = table.rows[0].cells[0]

    # Fill in the username...
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('username', qct.header['username'])

    # Fill in the date...
    cell.paragraphs[1].text = cell.paragraphs[1].text.replace('testdate', qct.header['testdate'])

    # On to the test steps and results table...
    table = doc.tables[2]
    
    # python-docx does not allow for search/replace at the table level, so we must
    # explicitly state which cell by column number we want to modify...
    for i in range(2, 12):
        key = table.columns[4].cells[i].text
        table.columns[4].cells[i].text = qct.results_text[key]
        table.columns[5].cells[i].text = 'X' if qct.results_pass[key] else ''
        table.columns[6].cells[i].text = '' if qct.results_pass[key] else 'X'

    # Complete the last table...
    table = doc.tables[4]
    table.rows[1].cells[0].text = ("Session log file: 3305-00101-%s-A.txt"
                                   % qct.header['formnumber'])

    # Set the document Title and Author attributes...
    doc.core_properties.title = ("SN_%s_QCT_Results_CTDMO-%s"
                                 % (qct.header['serialnumber'], qct.header['seriesletter']))
    author_inits = qct.header['username'].split()
    author_inits[:-1] = [init[0] + '.' for init in author_inits[:-1]]
    doc.core_properties.author = ' '.join(author_inits)

    # Save a copy of the results file with the form number in the title...
    doc.save('save/3305-00101-%s.docx' % formnumber)
